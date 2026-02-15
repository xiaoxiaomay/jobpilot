"""
DOCX Builder Module.
Creates ATS-compatible .docx files for resumes and cover letters using python-docx.
"""

import os
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")


def _set_page_margins(doc, margin_inches=0.5):
    """Set page margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def _set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_horizontal_line(doc):
    """Add a thin horizontal line separator."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_resume(
    content: dict,
    personal: dict,
    education: list,
) -> BytesIO:
    """
    Build a resume .docx from structured content.

    Args:
        content: Generated resume content with keys:
            summary, experiences, skills
        personal: Personal info dict (name, email, phone, location, linkedin, github)
        education: List of education dicts

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()
    _set_page_margins(doc)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Name ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(name_p, after=2)
    name_run = name_p.add_run(personal.get("name", ""))
    name_run.font.size = Pt(14)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # --- Contact Info ---
    contact_parts = []
    if personal.get("location"):
        contact_parts.append(personal["location"])
    if personal.get("phone"):
        contact_parts.append(personal["phone"])
    if personal.get("email"):
        contact_parts.append(personal["email"])
    if personal.get("linkedin"):
        contact_parts.append(personal["linkedin"])
    if personal.get("github"):
        contact_parts.append(personal["github"])

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(contact_p, after=4)
    contact_run = contact_p.add_run(" | ".join(contact_parts))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _add_horizontal_line(doc)

    # --- Professional Summary ---
    summary_header = doc.add_paragraph()
    _set_paragraph_spacing(summary_header, before=4, after=2)
    run = summary_header.add_run("PROFESSIONAL SUMMARY")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    summary_p = doc.add_paragraph()
    _set_paragraph_spacing(summary_p, after=4)
    summary_run = summary_p.add_run(content.get("summary", ""))
    summary_run.font.size = Pt(10)

    _add_horizontal_line(doc)

    # --- Technical Skills ---
    skills_header = doc.add_paragraph()
    _set_paragraph_spacing(skills_header, before=4, after=2)
    run = skills_header.add_run("TECHNICAL SKILLS")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    skills = content.get("skills", {})
    skill_labels = {
        "programming": "Programming",
        "ml_frameworks": "ML/AI Frameworks",
        "data_tools": "Data Tools",
        "methods": "Methods & Techniques",
        "ml_methods": "Methods & Techniques",
        "ml_techniques": "Methods & Techniques",
        "cloud": "Cloud & Infrastructure",
        "soft_skills": "Professional Skills",
        "tools_platforms": "Tools & Platforms",
    }
    for cat, skill_list in skills.items():
        if not skill_list:
            continue
        label = skill_labels.get(cat, cat.replace("_", " ").title())
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=1)
        bold_run = p.add_run(f"{label}: ")
        bold_run.font.bold = True
        bold_run.font.size = Pt(10)
        skill_run = p.add_run(", ".join(skill_list))
        skill_run.font.size = Pt(10)

    _add_horizontal_line(doc)

    # --- Professional Experience ---
    exp_header = doc.add_paragraph()
    _set_paragraph_spacing(exp_header, before=4, after=2)
    run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    for exp in content.get("experiences", []):
        # Title + Company line
        title_p = doc.add_paragraph()
        _set_paragraph_spacing(title_p, before=4, after=1)

        title_run = title_p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
        title_run.font.bold = True
        title_run.font.size = Pt(10)

        # Dates + Location line
        dates_p = doc.add_paragraph()
        _set_paragraph_spacing(dates_p, after=2)
        location = exp.get("location", "")
        dates = exp.get("dates", "")
        dates_text = f"{location}  |  {dates}" if location else dates
        dates_run = dates_p.add_run(dates_text)
        dates_run.font.size = Pt(9)
        dates_run.font.italic = True
        dates_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Bullet points
        for bullet in exp.get("bullets", []):
            bullet_p = doc.add_paragraph(style='List Bullet')
            _set_paragraph_spacing(bullet_p, after=1)
            bullet_run = bullet_p.add_run(bullet)
            bullet_run.font.size = Pt(10)

    _add_horizontal_line(doc)

    # --- Education ---
    edu_header = doc.add_paragraph()
    _set_paragraph_spacing(edu_header, before=4, after=2)
    run = edu_header.add_run("EDUCATION")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    for edu in education:
        edu_p = doc.add_paragraph()
        _set_paragraph_spacing(edu_p, after=1)

        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        dates = edu.get("dates", "")

        deg_run = edu_p.add_run(f"{degree}")
        deg_run.font.bold = True
        deg_run.font.size = Pt(10)

        inst_run = edu_p.add_run(f" — {institution}")
        inst_run.font.size = Pt(10)

        if dates:
            dates_run = edu_p.add_run(f"  ({dates})")
            dates_run.font.size = Pt(9)
            dates_run.font.italic = True
            dates_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if edu.get("second_degree"):
            sd_p = doc.add_paragraph()
            _set_paragraph_spacing(sd_p, after=1)
            sd_run = sd_p.add_run(f"{edu['second_degree']}")
            sd_run.font.bold = True
            sd_run.font.size = Pt(10)
            sd_run2 = sd_p.add_run(f" — {institution}")
            sd_run2.font.size = Pt(10)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_cover_letter(
    content: dict,
    personal: dict,
    company: str,
    title: str,
    hiring_manager: str = "Hiring Manager",
) -> BytesIO:
    """
    Build a cover letter .docx from structured content.

    Args:
        content: Cover letter content with keys:
            opening, body_paragraph_1, body_paragraph_2, body_paragraph_3, closing
        personal: Personal info dict.
        company: Company name.
        title: Job title.
        hiring_manager: Name of hiring manager (default: "Hiring Manager").

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()
    _set_page_margins(doc)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Date ---
    date_p = doc.add_paragraph()
    _set_paragraph_spacing(date_p, after=8)
    date_run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)

    # --- Sender info ---
    sender_p = doc.add_paragraph()
    _set_paragraph_spacing(sender_p, after=4)
    sender_lines = [
        personal.get("name", ""),
        personal.get("location", ""),
        personal.get("email", ""),
        personal.get("phone", ""),
    ]
    sender_run = sender_p.add_run("\n".join(line for line in sender_lines if line))
    sender_run.font.size = Pt(10)

    # --- Recipient ---
    recipient_p = doc.add_paragraph()
    _set_paragraph_spacing(recipient_p, after=8)
    recipient_run = recipient_p.add_run(
        f"Dear {hiring_manager},\n"
        f"Re: {title} at {company}"
    )
    recipient_run.font.size = Pt(10)

    # --- Body paragraphs ---
    cl = content if isinstance(content, dict) else {}

    for key in ["opening", "body_paragraph_1", "body_paragraph_2", "body_paragraph_3"]:
        text = cl.get(key, "")
        if not text:
            continue
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=6)
        run = p.add_run(text)
        run.font.size = Pt(10)

    # --- Closing ---
    closing_text = cl.get("closing", f"Sincerely,\n{personal.get('name', '')}")
    closing_p = doc.add_paragraph()
    _set_paragraph_spacing(closing_p, before=8)
    closing_run = closing_p.add_run(closing_text)
    closing_run.font.size = Pt(10)

    # Contact footer
    footer_p = doc.add_paragraph()
    _set_paragraph_spacing(footer_p, before=12)
    footer_parts = []
    if personal.get("linkedin"):
        footer_parts.append(personal["linkedin"])
    if personal.get("github"):
        footer_parts.append(personal["github"])
    if footer_parts:
        footer_run = footer_p.add_run(" | ".join(footer_parts))
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def save_docx(buffer: BytesIO, filename: str, output_dir: str = OUTPUT_DIR) -> str:
    """Save a BytesIO buffer to a .docx file and return the path."""
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, filename)
    with open(path, "wb") as f:
        f.write(buffer.getvalue())
    return path
